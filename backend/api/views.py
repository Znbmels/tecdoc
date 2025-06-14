from rest_framework import generics, permissions, status, serializers
from rest_framework_simplejwt.views import TokenObtainPairView
from rest_framework_simplejwt.serializers import TokenObtainPairSerializer
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from django.core.mail import send_mail
from django.conf import settings
from .models import User, Document, DocumentCollaborator, ShareToken
from .serializers import UserSerializer, DocumentSerializer, DocumentCollaboratorSerializer, ShareDocumentSerializer, ShareTokenSerializer
from django.http import HttpResponse, JsonResponse
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from docx import Document as DocxDocument
import io

# Регистрация пользователя
@api_view(['POST'])
def register(request):
    serializer = UserSerializer(data=request.data)
    if serializer.is_valid():
        user = serializer.save()
        return Response({"message": "User registered", "user_id": user.id}, status=status.HTTP_201_CREATED)
    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

# Кастомный сериализатор для логина с email
class CustomTokenObtainPairSerializer(TokenObtainPairSerializer):
    def validate(self, attrs):
        print("Received credentials:", attrs)
        email = attrs.get('email')
        user = User.objects.filter(email=email).first()
        if user:
            print(f"Found user: {user.username}")
            attrs['username'] = user.username
        else:
            print(f"User with email {email} not found")
            raise serializers.ValidationError({"email": "No active account found with the given email"})
        return super().validate(attrs)

# Кастомное представление для логина
class CustomTokenObtainPairView(TokenObtainPairView):
    serializer_class = CustomTokenObtainPairSerializer

# Список документов и создание нового
class DocumentListCreate(generics.ListCreateAPIView):
    from .serializers import DocumentSerializer
    serializer_class = DocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return (Document.objects.filter(owner=user) | Document.objects.filter(collaborators__user=user)).distinct()

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

# Просмотр, редактирование и удаление документа
class DocumentDetail(generics.RetrieveUpdateDestroyAPIView):
    serializer_class = DocumentSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        user = self.request.user
        return (Document.objects.filter(owner=user) | Document.objects.filter(collaborators__user=user)).distinct()

    def update(self, request, *args, **kwargs):
        document = self.get_object()
        user = request.user

        # Проверяем права на редактирование
        is_owner = document.owner == user
        is_edit_collaborator = document.collaborators.filter(user=user, role='edit').exists()

        if not is_owner and not is_edit_collaborator:
            return Response({"error": "You do not have permission to edit this document."}, status=status.HTTP_403_FORBIDDEN)

        return super().update(request, *args, **kwargs)

    def destroy(self, request, *args, **kwargs):
        document = self.get_object()
        user = request.user

        # Только владелец может удалить документ
        if document.owner != user:
            return Response({"error": "Only the owner can delete this document."}, status=status.HTTP_403_FORBIDDEN)

        return super().destroy(request, *args, **kwargs)

# Приглашение collaborator
@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def invite_collaborator(request, id):
    from .serializers import DocumentCollaboratorSerializer
    try:
        document = Document.objects.get(id=id, owner=request.user)
        email = request.data.get('email')
        role = request.data.get('role', 'view')
        user = User.objects.get(email=email)
        collaborator, created = DocumentCollaborator.objects.get_or_create(
            user=user,
            document=document,
            defaults={'role': role}
        )
        send_mail(
            'Invitation to collaborate',
            f'You have been invited to collaborate on document "{document.title}" with {role} access.',
            settings.EMAIL_HOST_USER,
            [email],
            fail_silently=False,
        )
        return Response({"message": "Invitation sent"}, status=status.HTTP_200_OK)
    except Document.DoesNotExist:
        return Response({"error": "Document not found or you are not the owner"}, status=status.HTTP_404_NOT_FOUND)
    except User.DoesNotExist:
        return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

@api_view(['POST'])
@permission_classes([permissions.IsAuthenticated])
def share_document(request):
    serializer = ShareDocumentSerializer(data=request.data)
    if serializer.is_valid():
        document_id = serializer.validated_data['document_id']
        email = serializer.validated_data['email']
        
        try:
            document = Document.objects.get(id=document_id, owner=request.user)
        except Document.DoesNotExist:
            return Response({"error": "Document not found or you are not the owner"}, status=status.HTTP_404_NOT_FOUND)

        try:
            user_to_share_with = User.objects.get(email=email)
        except User.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

        # Создаем коллаборатора с правами просмотра
        collaborator, created = DocumentCollaborator.objects.get_or_create(
            user=user_to_share_with,
            document=document,
            defaults={'role': 'view'}
        )

        share_token = ShareToken.objects.create(document=document, email=email)

        # Отправляем email уведомление
        share_link = f"http://localhost:8081/shared-document/{share_token.token}"
        send_mail(
            'Document Shared with You',
            f'You have been given access to the document "{document.title}". You can view it in your documents list or using this link: {share_link}',
            settings.EMAIL_HOST_USER,
            [email],
            fail_silently=False,
        )
        
        return Response({
            "message": "Share link created successfully.", 
            "token": share_token.token
        }, status=status.HTTP_201_CREATED)

    return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@api_view(['GET'])
def retrieve_shared_document(request, token):
    try:
        share_token = ShareToken.objects.get(token=token, is_active=True)
        document = share_token.document
        serializer = DocumentSerializer(document)
        return Response(serializer.data, status=status.HTTP_200_OK)
    except ShareToken.DoesNotExist:
        return Response({"error": "Invalid or expired token"}, status=status.HTTP_404_NOT_FOUND)

def download_document(request, pk):
    print(f"=== DOWNLOAD REQUEST START ===")
    print(f"Document ID: {pk}")
    print(f"Format: {request.GET.get('format', 'pdf')}")
    print(f"User: {request.user}")
    print(f"Request path: {request.path}")
    print(f"Authorization header: {request.META.get('HTTP_AUTHORIZATION', 'None')}")

    # Проверка аутентификации JWT
    from rest_framework_simplejwt.authentication import JWTAuthentication
    from rest_framework.request import Request

    # Создаем DRF Request для аутентификации
    drf_request = Request(request)
    jwt_auth = JWTAuthentication()

    try:
        user_auth_tuple = jwt_auth.authenticate(drf_request)
        if user_auth_tuple is None:
            return JsonResponse({"error": "Authentication required"}, status=401)

        user, token = user_auth_tuple
        print(f"Authenticated user: {user}")
    except Exception as e:
        print(f"Authentication error: {e}")
        return JsonResponse({"error": "Invalid authentication"}, status=401)

    try:
        document = Document.objects.get(pk=pk)
        print(f"Document found: {document.title}")
        print(f"Document owner: {document.owner}")
        
        # Check if the user is the owner or a collaborator
        is_owner = document.owner == user
        is_collaborator = document.collaborators.filter(user=user).exists()
        print(f"User permissions: is_owner={is_owner}, is_collaborator={is_collaborator}")

        if not is_owner and not is_collaborator:
            print("User does not have permission to access this document.")
            return JsonResponse({"error": "You do not have permission to access this document."}, status=403)

        file_format = request.GET.get('format', 'pdf').lower()
        print(f"Generating {file_format} file for document: {document.title}")

        if file_format == 'pdf':
            # Создаем PDF документ
            buffer = io.BytesIO()
            p = canvas.Canvas(buffer, pagesize=letter)
            p.drawString(72, 800, document.title)
            text = p.beginText(72, 780)
            text.textLines(document.content)
            p.drawText(text)
            p.showPage()
            p.save()
            buffer.seek(0)
            
            # Устанавливаем заголовки для загрузки файла
            response = HttpResponse(buffer.getvalue(), content_type='application/pdf')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.pdf"'
            print("PDF generated successfully")
            return response

        elif file_format == 'docx':
            # Создаем DOCX документ
            buffer = io.BytesIO()
            doc = DocxDocument()
            doc.add_heading(document.title, level=1)
            doc.add_paragraph(document.content)
            doc.save(buffer)
            buffer.seek(0)
            
            # Устанавливаем заголовки для загрузки файла
            response = HttpResponse(buffer.getvalue(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{document.title}.docx"'
            print("DOCX generated successfully")
            return response

        else:
            print(f"Unsupported format: {file_format}")
            return JsonResponse({"error": f"Unsupported format: {file_format}. Please choose 'pdf' or 'docx'."}, status=400)

    except Document.DoesNotExist:
        print(f"Document with id {pk} not found")
        return JsonResponse({"error": "Document not found"}, status=404)
    except Exception as e:
        print(f"Error while generating document: {str(e)}")
        return JsonResponse({"error": f"Error: {str(e)}"}, status=500)